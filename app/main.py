from __future__ import annotations

import io
import os
import re
import secrets
import time
import urllib.parse
import uuid
from datetime import datetime, timezone
from pathlib import Path

import requests
from argon2 import PasswordHasher
from argon2.exceptions import VerifyMismatchError
from docx import Document
from fastapi import FastAPI, File, Form, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from starlette.middleware.sessions import SessionMiddleware

from app.services import elite_companies, infer_roles, ingest_jobs_for_companies, rank_jobs, tailor_resume
from app.storage import DataStore

BASE_DIR = Path(__file__).resolve().parent.parent
UPLOAD_DIR = BASE_DIR / "data" / "uploads"
EXPORT_DIR = BASE_DIR / "data" / "exports"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

SESSION_SECRET = os.getenv("SESSION_SECRET", "dev-secret-change-me")
COOKIE_SECURE = os.getenv("COOKIE_SECURE", "false").lower() == "true"
ph = PasswordHasher()

LINKEDIN_CLIENT_ID = os.getenv("LINKEDIN_CLIENT_ID", "")
LINKEDIN_CLIENT_SECRET = os.getenv("LINKEDIN_CLIENT_SECRET", "")
LINKEDIN_REDIRECT_URI = os.getenv("LINKEDIN_REDIRECT_URI", "http://localhost:8787/auth/linkedin/callback")

app = FastAPI(title="Mage-Luz MVP")
app.add_middleware(SessionMiddleware, secret_key=SESSION_SECRET, same_site="lax", https_only=COOKIE_SECURE, max_age=60 * 60 * 24)
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "app" / "static")), name="static")
templates = Jinja2Templates(directory=str(BASE_DIR / "app" / "templates"))
store = DataStore(BASE_DIR / "data" / "db.json")


def _extract_text(file_name: str, data: bytes) -> str:
    lower = file_name.lower()
    if lower.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    if lower.endswith(".pdf"):
        from pypdf import PdfReader

        pdf = PdfReader(io.BytesIO(data))
        return "\n".join([p.extract_text() or "" for p in pdf.pages])
    if lower.endswith(".docx"):
        import docx

        doc = docx.Document(io.BytesIO(data))
        return "\n".join([p.text for p in doc.paragraphs])
    return data.decode("utf-8", errors="ignore")


def _csrf_token(request: Request) -> str:
    token = request.session.get("csrf_token")
    if not token:
        token = secrets.token_urlsafe(32)
        request.session["csrf_token"] = token
    return token


def _check_csrf(request: Request, token: str) -> bool:
    expected = request.session.get("csrf_token")
    return bool(expected and token and secrets.compare_digest(expected, token))


def _current_user(request: Request, state: dict):
    user_id = request.session.get("user_id")
    if not user_id:
        return None
    return next((u for u in state.get("users", []) if u["id"] == user_id), None)


def _require_user(request: Request, state: dict):
    user = _current_user(request, state)
    if not user:
        return None, RedirectResponse(url="/login", status_code=303)
    return user, None


def _render_login(request: Request, error: str | None = None):
    return templates.TemplateResponse(
        "login.html",
        {
            "request": request,
            "error": error,
            "csrf_token": _csrf_token(request),
            "linkedin_ready": bool(LINKEDIN_CLIENT_ID and LINKEDIN_CLIENT_SECRET and LINKEDIN_REDIRECT_URI),
        },
    )


def _rate_limited(state: dict, key: str) -> tuple[bool, int]:
    rec = (state.get("login_attempts") or {}).get(key, {"count": 0, "locked_until": 0})
    now = int(time.time())
    if rec.get("locked_until", 0) > now:
        return True, int(rec["locked_until"] - now)
    return False, 0


def _fail_login(state: dict, key: str):
    attempts = state.setdefault("login_attempts", {})
    rec = attempts.get(key, {"count": 0, "locked_until": 0})
    rec["count"] += 1
    if rec["count"] >= 5:
        rec["locked_until"] = int(time.time()) + 15 * 60
        rec["count"] = 0
    attempts[key] = rec


@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    return _render_login(request)


@app.get("/auth/linkedin/start")
def linkedin_start(request: Request):
    if not (LINKEDIN_CLIENT_ID and LINKEDIN_CLIENT_SECRET and LINKEDIN_REDIRECT_URI):
        return _render_login(request, "LinkedIn OAuth is not configured.")
    state_token = secrets.token_urlsafe(24)
    request.session["linkedin_oauth_state"] = state_token
    q = {
        "response_type": "code",
        "client_id": LINKEDIN_CLIENT_ID,
        "redirect_uri": LINKEDIN_REDIRECT_URI,
        "state": state_token,
        "scope": "openid profile email",
    }
    return RedirectResponse("https://www.linkedin.com/oauth/v2/authorization?" + urllib.parse.urlencode(q), status_code=302)


@app.get("/auth/linkedin/callback")
def linkedin_callback(request: Request, code: str = "", state: str = ""):
    expected = request.session.get("linkedin_oauth_state")
    if not expected or not state or not secrets.compare_digest(expected, state):
        return _render_login(request, "LinkedIn login failed.")
    try:
        tok = requests.post(
            "https://www.linkedin.com/oauth/v2/accessToken",
            data={
                "grant_type": "authorization_code",
                "code": code,
                "redirect_uri": LINKEDIN_REDIRECT_URI,
                "client_id": LINKEDIN_CLIENT_ID,
                "client_secret": LINKEDIN_CLIENT_SECRET,
            },
            timeout=20,
        )
        tok.raise_for_status()
        access = tok.json().get("access_token")
        prof = requests.get("https://api.linkedin.com/v2/userinfo", headers={"Authorization": f"Bearer {access}"}, timeout=20)
        prof.raise_for_status()
        p = prof.json()
    except Exception:
        return _render_login(request, "LinkedIn token/profile exchange failed.")

    sub = str(p.get("sub") or p.get("id") or "")
    email = p.get("email") or f"linkedin_{sub}@local"
    db = store.load()
    user = next((u for u in db.get("users", []) if u.get("linkedin_sub") == sub), None) or next(
        (u for u in db.get("users", []) if u.get("email", "").lower() == email.lower()), None
    )
    if not user:
        user = {
            "id": str(uuid.uuid4()),
            "email": email,
            "password_hash": "",
            "linkedin_sub": sub,
            "is_paid": False,
            "plan": "free",
            "preferred_locations": ["New York", "San Francisco"],
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        db.setdefault("users", []).append(user)
    else:
        user["linkedin_sub"] = sub
    store.save(db)
    request.session["user_id"] = user["id"]
    return RedirectResponse(url="/", status_code=303)


@app.post("/register")
def register(request: Request, csrf_token: str = Form(...), email: str = Form(...), password: str = Form(...)):
    if not _check_csrf(request, csrf_token):
        return _render_login(request, "Refresh and try again.")
    db = store.load()
    if any(u["email"].lower() == email.lower() for u in db.get("users", [])):
        return _render_login(request, "Email already exists.")
    user = {
        "id": str(uuid.uuid4()),
        "email": email,
        "password_hash": ph.hash(password),
        "linkedin_sub": "",
        "is_paid": False,
        "plan": "free",
        "preferred_locations": ["New York", "San Francisco"],
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    db.setdefault("users", []).append(user)
    store.save(db)
    request.session["user_id"] = user["id"]
    return RedirectResponse(url="/", status_code=303)


@app.post("/login")
def login(request: Request, csrf_token: str = Form(...), email: str = Form(...), password: str = Form(...)):
    if not _check_csrf(request, csrf_token):
        return _render_login(request, "Refresh and try again.")
    db = store.load()
    key = f"{email.lower()}::{(request.client.host if request.client else 'unknown')}"
    limited, wait = _rate_limited(db, key)
    if limited:
        return _render_login(request, f"Too many attempts. Try again in {wait}s.")
    user = next((u for u in db.get("users", []) if u.get("email", "").lower() == email.lower()), None)
    if not user:
        _fail_login(db, key)
        store.save(db)
        return _render_login(request, "Invalid credentials.")
    ok = False
    if user.get("password_hash"):
        try:
            ok = ph.verify(user["password_hash"], password)
        except VerifyMismatchError:
            ok = False
    if not ok:
        _fail_login(db, key)
        store.save(db)
        return _render_login(request, "Invalid credentials.")
    db.setdefault("login_attempts", {}).pop(key, None)
    store.save(db)
    request.session["user_id"] = user["id"]
    return RedirectResponse(url="/", status_code=303)


@app.post("/logout")
def logout(request: Request, csrf_token: str = Form(...)):
    if _check_csrf(request, csrf_token):
        request.session.clear()
    return RedirectResponse(url="/login", status_code=303)


@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    db = store.load()
    user, redirect = _require_user(request, db)
    if redirect:
        return redirect

    resumes = [r for r in db.get("resumes", []) if r.get("user_id") == user["id"]]
    matches = [m for m in db.get("matches", []) if m.get("user_id") == user["id"]]
    variants = [v for v in db.get("variants", []) if v.get("user_id") == user["id"]]

    visible = matches if user.get("is_paid") else matches[:3]
    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
            "user": user,
            "has_data": bool(matches),
            "resumes": resumes,
            "matches": visible,
            "locked_count": max(0, len(matches) - len(visible)),
            "variants": variants,
            "csrf_token": _csrf_token(request),
        },
    )


@app.post("/start")
async def start_flow(request: Request, csrf_token: str = Form(...), preferred_locations: str = Form(...), file: UploadFile = File(...)):
    if not _check_csrf(request, csrf_token):
        return RedirectResponse(url="/", status_code=303)

    db = store.load()
    user, redirect = _require_user(request, db)
    if redirect:
        return redirect

    user["preferred_locations"] = [x.strip() for x in preferred_locations.split(",") if x.strip()]

    file_bytes = await file.read()
    resume_id = str(uuid.uuid4())
    safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", file.filename or "resume")
    out_path = UPLOAD_DIR / f"{resume_id}_{safe_name}"
    out_path.write_bytes(file_bytes)
    parsed_text = _extract_text(safe_name, file_bytes)
    role_profile = infer_roles(parsed_text)

    resume = {
        "id": resume_id,
        "user_id": user["id"],
        "file_name": safe_name,
        "file_path": str(out_path),
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
        "parsed_text": parsed_text[:25000],
        "role_profile": role_profile,
    }
    db["resumes"] = [r for r in db.get("resumes", []) if r.get("user_id") != user["id"]] + [resume]

    companies = elite_companies()
    jobs = ingest_jobs_for_companies(companies, role_profile)
    matches = rank_jobs(resume, jobs, companies, user.get("preferred_locations", []))
    for m in matches:
        m["user_id"] = user["id"]

    db["jobs"] = jobs
    db["matches"] = [m for m in db.get("matches", []) if m.get("user_id") != user["id"]] + matches
    db["variants"] = [v for v in db.get("variants", []) if v.get("user_id") != user["id"]]
    store.save(db)
    return RedirectResponse(url="/", status_code=303)


@app.post("/mock-upgrade")
def mock_upgrade(request: Request, csrf_token: str = Form(...)):
    if not _check_csrf(request, csrf_token):
        return RedirectResponse(url="/", status_code=303)
    db = store.load()
    user, redirect = _require_user(request, db)
    if redirect:
        return redirect
    user["is_paid"] = True
    user["plan"] = "pro_7_99"
    store.save(db)
    return RedirectResponse(url="/", status_code=303)


@app.post("/tailor/{match_id}")
def tailor(request: Request, match_id: str, csrf_token: str = Form(...)):
    if not _check_csrf(request, csrf_token):
        return RedirectResponse(url="/", status_code=303)
    db = store.load()
    user, redirect = _require_user(request, db)
    if redirect:
        return redirect

    resume = next((r for r in db.get("resumes", []) if r.get("user_id") == user["id"]), None)
    match = next((m for m in db.get("matches", []) if m.get("id") == match_id and m.get("user_id") == user["id"]), None)
    if not resume or not match:
        return RedirectResponse(url="/", status_code=303)

    job = next((j for j in db.get("jobs", []) if j.get("id") == match.get("job_id")), None)
    if not job:
        return RedirectResponse(url="/", status_code=303)

    variant = tailor_resume(resume, job, match)
    variant["user_id"] = user["id"]
    db["variants"] = [v for v in db.get("variants", []) if not (v.get("user_id") == user["id"] and v.get("job_id") == variant["job_id"])] + [variant]
    store.save(db)
    return RedirectResponse(url="/", status_code=303)


@app.get("/export/{variant_id}.{fmt}")
def export_variant(request: Request, variant_id: str, fmt: str):
    db = store.load()
    user, redirect = _require_user(request, db)
    if redirect:
        return redirect

    variant = next((v for v in db.get("variants", []) if v.get("id") == variant_id and v.get("user_id") == user["id"]), None)
    if not variant:
        return RedirectResponse(url="/", status_code=303)

    safe_base = re.sub(r"[^a-zA-Z0-9_-]", "_", f"{variant['company_name']}_{variant['job_title']}_{variant_id[:8]}")

    if fmt == "docx":
        p = EXPORT_DIR / f"{safe_base}.docx"
        doc = Document()
        doc.add_heading("Tailored Resume Pack", level=1)
        doc.add_paragraph(variant.get("summary", ""))
        doc.add_paragraph(f"Role link: {variant.get('apply_url', '')}")
        doc.add_paragraph(variant.get("variant_text", ""))
        doc.save(str(p))
        return FileResponse(str(p), filename=p.name)

    if fmt == "pdf":
        p = EXPORT_DIR / f"{safe_base}.pdf"
        c = canvas.Canvas(str(p), pagesize=LETTER)
        y = 760
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y, "Tailored Resume Pack")
        y -= 28
        c.setFont("Helvetica", 11)
        for line in [variant.get("summary", ""), f"Role link: {variant.get('apply_url', '')}", "", variant.get("variant_text", "")]:
            for seg in str(line).split("\n"):
                c.drawString(50, y, seg[:110])
                y -= 16
                if y < 60:
                    c.showPage()
                    y = 760
        c.save()
        return FileResponse(str(p), filename=p.name)

    return RedirectResponse(url="/", status_code=303)
